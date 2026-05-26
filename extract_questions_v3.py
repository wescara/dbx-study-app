from docx import Document
import pandas as pd
import re
from datetime import datetime

doc_path = 'c:\\dbx-study-app\\data\\DCDEA Practice Exam 2.docx'
doc = Document(doc_path)

questions_data = []
i = 0

while i < len(doc.paragraphs):
    para = doc.paragraphs[i]
    text = para.text.strip()
    
    # Look for question start
    if text.startswith('Question '):
        match = re.match(r'Question (\d+):', text)
        if match:
            q_num = int(match.group(1))
            current_question = {
                'Number': q_num,
                'Question': '',
                'A': '',
                'B': '',
                'C': '',
                'D': '',
                'E': '',
                'F': '',
                'Explanation': '',
                'CorrectAnswer': ''
            }
            
            # Next line is the main question text
            i += 1
            if i < len(doc.paragraphs):
                main_q = doc.paragraphs[i].text.strip()
            
            # Next line is usually the specific question
            i += 1
            if i < len(doc.paragraphs):
                specific_q = doc.paragraphs[i].text.strip()
                current_question['Question'] = f"{main_q} {specific_q}"
            
            # Next lines are options (until we hit separator)
            options = []
            i += 1
            while i < len(doc.paragraphs):
                opt_text = doc.paragraphs[i].text.strip()
                
                # Stop at separator
                if '---' in opt_text or opt_text.lower().startswith('overall explanation') or opt_text.lower().startswith('documentation'):
                    break
                
                if opt_text:
                    options.append(opt_text)
                
                i += 1
            
            # Assign options to A, B, C, D, E, F
            for j, opt in enumerate(options):
                if j < 6:
                    current_question[chr(65 + j)] = opt
            
            # Now look for explanation, documentation, and correct answer
            while i < len(doc.paragraphs):
                para_text = doc.paragraphs[i].text.strip()
                
                if para_text.lower().startswith('correct answer'):
                    match = re.search(r'Correct Answer[:\s]*(.+)', para_text, re.IGNORECASE)
                    if match:
                        current_question['CorrectAnswer'] = match.group(1).strip()
                    break
                
                if para_text.lower().startswith('overall explanation'):
                    # Collect explanation until separator or documentation
                    i += 1
                    expl = []
                    while i < len(doc.paragraphs):
                        exp_text = doc.paragraphs[i].text.strip()
                        if '---' in exp_text or exp_text.lower().startswith('documentation'):
                            break
                        if exp_text and not exp_text.lower().startswith('overall explanation'):
                            expl.append(exp_text)
                        i += 1
                    current_question['Explanation'] = ' '.join(expl)
                    continue
                
                i += 1
            
            questions_data.append(current_question)
    
    i += 1

print(f"Extracted {len(questions_data)} questions")
print(f"\nFirst 2 questions:")
for q in questions_data[:2]:
    print(f"\nQ{q['Number']}:")
    print(f"  Question: {q['Question'][:100]}")
    print(f"  A: {q['A'][:80]}")
    print(f"  B: {q['B'][:80]}")
    print(f"  C: {q['C'][:80]}")
    print(f"  D: {q['D'][:80]}")
    print(f"  Correct: {q['CorrectAnswer']}")

# Create DataFrame
df_new = pd.DataFrame()
df_new['Source'] = 'DCDEA Practice Exam 2'
df_new['QID'] = [f"DCDEA2_{i:03d}" for i in range(1, len(questions_data) + 1)]
df_new['Exam'] = 2
df_new['Number'] = [q['Number'] for q in questions_data]
df_new['Topic'] = ''
df_new['Subtopic'] = ''
df_new['Difficulty'] = ''
df_new['Question'] = [q['Question'] for q in questions_data]
df_new['A'] = [q['A'] for q in questions_data]
df_new['B'] = [q['B'] for q in questions_data]
df_new['C'] = [q['C'] for q in questions_data]
df_new['D'] = [q['D'] for q in questions_data]
df_new['E'] = [q['E'] for q in questions_data]
df_new['F'] = [q['F'] for q in questions_data]
df_new['CorrectLetter'] = [q['CorrectAnswer'] for q in questions_data]
df_new['Explanation'] = [q['Explanation'] for q in questions_data]
df_new['Notes'] = ''
df_new['LastReviewed'] = ''
df_new['CorrectCount'] = ''
df_new['IncorrectCount'] = ''
df_new['FlagForReview'] = ''
df_new['Tags'] = ''
df_new['Graphic'] = ''
df_new['Deeper Dive'] = ''
df_new['VerifiedAnswer'] = ''
df_new['VerificationStatus'] = ''
df_new['AnswerMatchesWorkbook'] = ''
df_new['VerificationNotes'] = ''
df_new['VerificationSourceKeys'] = ''
df_new['VerifiedOn'] = ''
df_new['ToDelete'] = ''
df_new['DateAdded'] = datetime.now()

# Save to Excel
output_file = 'c:\\dbx-study-app\\data\\DCDEA_Practice_Exam_2_NEW_QUESTIONS.xlsx'
df_new.to_excel(output_file, index=False, sheet_name='Questions')

print(f"\n✓ Created Excel file: {output_file}")
print(f"  Total questions: {len(df_new)}")
