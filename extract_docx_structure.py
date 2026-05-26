from docx import Document
import os

docx_path = 'data/DCDEA Practice Exam 2.docx'

# Load the document
doc = Document(docx_path)

print("Document structure analysis:")
print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")
print()

# Show first 100 paragraphs to understand structure
print("First 100 paragraphs content:")
for i, para in enumerate(doc.paragraphs[:100]):
    text = para.text.strip()
    if text:
        print(f"P{i}: {text[:120]}")
print()

# Check for images/shapes
print(f"Inline shapes in document: {len(doc.inline_shapes)}")
print()

# List all tables if any
if doc.tables:
    print(f"Tables found: {len(doc.tables)}")
    for t_idx, table in enumerate(doc.tables[:3]):
        print(f"\nTable {t_idx}: {len(table.rows)} rows x {len(table.columns)} cols")
        for r_idx, row in enumerate(table.rows[:5]):
            print(f"  Row {r_idx}: {[cell.text[:60] for cell in row.cells]}")
