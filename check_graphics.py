from docx import Document
from docx.oxml import parse_xml
import os

doc = Document('c:\\dbx-study-app\\data\\DCDEA Practice Exam 2.docx')

print('Checking for graphics and embedded content...')
print(f'Total paragraphs: {len(doc.paragraphs)}')
print(f'Total tables: {len(doc.tables)}')

# Check for inline shapes/images
shape_count = 0
for rel in doc.part.rels.values():
    if 'image' in rel.reltype:
        shape_count += 1
        print(f'Found image: {rel.target_ref}')

print(f'Total images found: {shape_count}')

# Check for drawing elements
drawing_count = 0
for para in doc.paragraphs:
    for run in para.runs:
        if hasattr(run, '_element'):
            drawings = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
            if drawings:
                drawing_count += len(drawings)
                print(f'Found drawing element in paragraph')

print(f'Total drawing elements: {drawing_count}')

# List all relationships
print('\nAll relationships in document:')
for rel_id, rel in doc.part.rels.items():
    reltype = rel.reltype.split("/")[-1]
    print(f'  {rel_id}: {reltype} -> {rel.target_ref}')
