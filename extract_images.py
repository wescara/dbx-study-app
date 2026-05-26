from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import re
import pandas as pd
from PIL import Image
from io import BytesIO
import os

# First, let's extract and inspect the images
docx_path = 'data/DCDEA Practice Exam 2.docx'
doc = Document(docx_path)

print("Extracting images from DOCX...")
output_dir = 'data/extracted_images_exam2'
os.makedirs(output_dir, exist_ok=True)

image_data = {}
for rel_id, rel in doc.part.rels.items():
    if "image" in rel.target_ref:
        image_name = rel.target_ref.split('/')[-1]
        image_part = rel.target_part
        image_bytes = image_part.blob
        
        # Save image to file
        image_path = os.path.join(output_dir, image_name)
        with open(image_path, 'wb') as f:
            f.write(image_bytes)
        
        image_data[rel_id] = {
            'name': image_name,
            'path': image_path,
            'size': len(image_bytes)
        }
        print(f"  Extracted {image_name} ({len(image_bytes)} bytes)")

print(f"\nTotal images extracted: {len(image_data)}")

# Now link images to questions by paragraph position
print("\nLinking images to questions...")
image_by_para = {}

for para_idx, para in enumerate(doc.paragraphs):
    # Check for drawing elements (images) in this paragraph
    for rel_id in doc.part.rels.keys():
        rel = doc.part.rels[rel_id]
        if "image" in rel.target_ref:
            # Check if this image appears in this paragraph
            if rel_id in str(para._element.xml):
                image_by_para[para_idx] = rel_id
                image_name = rel.target_ref.split('/')[-1]
                para_text = para.text[:50] if para.text else "(empty)"
                print(f"  Para {para_idx}: {image_name} - {para_text}...")

print(f"\nImages found in {len(image_by_para)} paragraphs")

# Save mapping for reference
with open('image_locations.txt', 'w') as f:
    for para_idx, rel_id in sorted(image_by_para.items()):
        f.write(f"Paragraph {para_idx}: {doc.part.rels[rel_id].target_ref}\n")
        if para_idx < len(doc.paragraphs):
            f.write(f"  Text: {doc.paragraphs[para_idx].text[:100]}\n\n")

print("Saved image location mapping to image_locations.txt")
